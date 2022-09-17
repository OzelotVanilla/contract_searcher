import docx
import re
from processor.statement_info import statement_dict
from util.console import console


def getSumTableRowPosition(doc) -> list[tuple[int, int]]: ...


sum_table_row_position = getSumTableRowPosition(docx.Document("D:/Code/contract_searcher/CoE Template.docx"))


def getSumTableRowPosition(doc) -> dict[int, tuple[int, int]]:
    """
    Pass the document, and return position of the row with "total score".
    return: (int table_index, int row_in_that_table_index)[]
    """
    positions = dict()
    for table_index in range(len(doc.tables)):
        for row_index in range(len(doc.tables[table_index].rows)):
            # If the first cell matches /^\d+\.?$/ after strip
            match_result = re.compile(
                "^(\\d+)\\.?$"
            ).match(doc.tables[table_index].rows[row_index].cells[0].text.strip())

            if match_result != None:
                positions[int(match_result.groups()[0])] = (table_index, row_index)

    return positions


def calculateSumedNumber(doc) -> dict[int, int]:
    """
    return { index_of_total: total_value }
    """
    total_sum = 0
    totals = dict()

    def isReachingNextTotalRow(table_index, row_index) -> bool:
        next_row_table_index, next_row_row_index = table_index, 0
        # If last row in one table
        if row_index == len(doc.tables[table_index].rows) - 1:
            # If reaching all tables end
            if table_index == len(doc.tables) - 1:
                return True
            next_row_table_index += 1
        else:
            next_row_row_index = row_index + 1

        return \
            re.compile(
                "^\\d+\\.?$"
            ).match(doc.tables[next_row_table_index].rows[next_row_row_index].cells[0].text.strip()) != None

    # Until touch the next "total" row, make sum
    current_total_index = 0
    table_len = len(doc.tables)
    for table_index in range(table_len):
        column_len = len(doc.tables[table_index].rows)
        for row_index in range(column_len):
            # Check if the second cell is empty, or number
            # If not, continue
            current_text = doc.tables[table_index].rows[row_index].cells[2].text.strip()
            if current_text == "1":
                total_sum += 1

            # Check if reach next "total" row
            if isReachingNextTotalRow(table_index, row_index):
                totals[current_total_index] = total_sum
                current_total_index += 1
                total_sum = 0
    del totals[0]
    return totals


def writeResultToWord(word_file_path: str, result_dict: dict[str, bool]) -> None:
    result_word_file = docx.Document(word_file_path)

    # Fill cell with result
    console.sublog("Filling cells", colour_rgb="b19a00", end="")
    for rule_name in result_dict.keys():
        table_index, row_index = statement_dict[rule_name].position
        is_rule_passed = result_dict[rule_name] == True
        cell_index = 2 if is_rule_passed else 3
        text_to_fill = "1" if is_rule_passed else "0"
        # Write 1 correspond to whether it has or not
        result_word_file.tables[table_index].rows[row_index].cells[cell_index].text = text_to_fill
    print(" ... done")

    # Fill the "sum" cell
    console.sublog("Filling sum cells", colour_rgb="b19a00", end="")
    fillSumValToCell(result_word_file)
    print(" ... done")

    # Save result
    result_word_file.save(word_file_path)


def fillSumValToCell(doc) -> None:
    # First calculate the sum-ed value
    sumed_value_dict = calculateSumedNumber(doc)

    # Fill it into corresponding cell
    global sum_table_row_position
    if sum_table_row_position == None:
        sum_table_row_position = getSumTableRowPosition(doc)
    for (sum_cell_index, sum_value) in sumed_value_dict.items():
        table_index, row_index = sum_table_row_position[sum_cell_index]
        doc.tables[table_index].rows[row_index].cells[5].text = str(sum_value)


def calculateAccuracy(standard_file_path: str, summoned_file_path: str):
    # TODO
    # Check with result got by manually check
    pass
